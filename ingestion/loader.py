from __future__ import annotations

import hashlib
import logging
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Callable, Dict, Optional, Tuple

from .models import Document
from .normalizer import NormalizationResult, TextNormalizer
from .text_utils import normalize_text

logger = logging.getLogger(__name__)

# Pre-compiled regex patterns for performance
_LIST_PATTERN = re.compile(r"^\s*(?:[-*+]|\d+\.)\s+")


class UnsupportedDocumentError(Exception):
    """Raised when the pipeline encounters an unsupported extension."""


def slugify(value: str) -> str:
    normalized = re.sub(r"[^a-zA-Z0-9]+", "-", value.lower()).strip("-")
    return normalized or "document"


def hash_file(path: Path) -> str:
    sha = hashlib.sha256()
    with path.open("rb") as file_obj:
        for chunk in iter(lambda: file_obj.read(8192), b""):
            sha.update(chunk)
    return sha.hexdigest()


class DocumentParseError(Exception):
    """Raised when a document cannot be parsed."""

    def __init__(self, message: str, path: Path, partial_content: str = ""):
        super().__init__(message)
        self.path = path
        self.partial_content = partial_content


def _parse_pdf_date(date_str: Optional[str]) -> Optional[str]:
    """Parse PDF date string to ISO format.

    PDF dates are typically in format: D:YYYYMMDDHHmmSS+HH'mm'
    """
    if not date_str:
        return None

    # Remove 'D:' prefix if present
    if date_str.startswith("D:"):
        date_str = date_str[2:]

    # Try to parse common PDF date formats
    try:
        # Basic format: YYYYMMDDHHMMSS
        if len(date_str) >= 14:
            year = date_str[0:4]
            month = date_str[4:6]
            day = date_str[6:8]
            hour = date_str[8:10]
            minute = date_str[10:12]
            second = date_str[12:14]
            return f"{year}-{month}-{day}T{hour}:{minute}:{second}"
        elif len(date_str) >= 8:
            year = date_str[0:4]
            month = date_str[4:6]
            day = date_str[6:8]
            return f"{year}-{month}-{day}"
    except (ValueError, IndexError):
        pass

    return None


def load_pdf(path: Path) -> Tuple[str, Dict[str, Any]]:
    """Load PDF file with error handling for corrupted/invalid files.

    Inserts page markers between pages in format [PAGE:N] where N is 1-indexed.
    This allows downstream chunking to track page boundaries.
    """
    from PyPDF2 import PdfReader
    from PyPDF2.errors import PdfReadError

    pages = []
    failed_pages = []
    metadata = {}

    try:
        reader = PdfReader(str(path))
        metadata["page_count"] = len(reader.pages)
        metadata["is_encrypted"] = reader.is_encrypted

        # Extract document metadata
        if reader.metadata:
            pdf_meta = reader.metadata
            if pdf_meta.title:
                metadata["title"] = pdf_meta.title
            if pdf_meta.author:
                metadata["author"] = pdf_meta.author
            if pdf_meta.creator:
                metadata["creator"] = pdf_meta.creator
            if pdf_meta.producer:
                metadata["producer"] = pdf_meta.producer
            if pdf_meta.subject:
                metadata["subject"] = pdf_meta.subject

            # Parse dates
            creation_date = _parse_pdf_date(pdf_meta.creation_date)
            if creation_date:
                metadata["creation_date"] = creation_date

            mod_date = _parse_pdf_date(pdf_meta.modification_date)
            if mod_date:
                metadata["modification_date"] = mod_date

        # Fallback: use filename as title if not in metadata
        if "title" not in metadata:
            metadata["title"] = path.stem

        if reader.is_encrypted:
            logger.warning(f"PDF is encrypted: {path}")
            metadata["parse_warning"] = "encrypted_pdf"
            return "", metadata

        for i, page in enumerate(reader.pages):
            try:
                text = page.extract_text() or ""
                page_text = text.strip()
                if page_text:
                    # Add page marker before content (1-indexed)
                    pages.append(f"[PAGE:{i + 1}]\n{page_text}")
                else:
                    # Empty page still gets marker for accurate page tracking
                    pages.append(f"[PAGE:{i + 1}]")
            except Exception as e:
                logger.warning(f"Failed to extract page {i + 1} from {path}: {e}")
                failed_pages.append(i + 1)
                pages.append(f"[PAGE:{i + 1}]")

        if failed_pages:
            metadata["failed_pages"] = failed_pages
            metadata["parse_warning"] = f"partial_extraction_{len(failed_pages)}_pages_failed"

    except PdfReadError as e:
        logger.error(f"Corrupted PDF file: {path} - {e}")
        raise DocumentParseError(f"Corrupted PDF: {e}", path)
    except Exception as e:
        logger.error(f"Failed to read PDF: {path} - {e}")
        raise DocumentParseError(f"PDF read error: {e}", path)

    return "\n\n".join(pages), metadata


def load_docx(path: Path) -> Tuple[str, Dict[str, Any]]:
    """Load DOCX file with table extraction and error handling."""
    import docx
    from docx.opc.exceptions import PackageNotFoundError
    import zipfile

    content_parts = []
    metadata = {}

    try:
        doc = docx.Document(str(path))

        # Extract core properties (document metadata)
        core_props = doc.core_properties
        if core_props.title:
            metadata["title"] = core_props.title
        if core_props.author:
            metadata["author"] = core_props.author
        if core_props.subject:
            metadata["subject"] = core_props.subject
        if core_props.keywords:
            metadata["keywords"] = core_props.keywords
        if core_props.category:
            metadata["category"] = core_props.category
        if core_props.comments:
            metadata["comments"] = core_props.comments
        if core_props.created:
            metadata["creation_date"] = core_props.created.isoformat()
        if core_props.modified:
            metadata["modification_date"] = core_props.modified.isoformat()
        if core_props.last_modified_by:
            metadata["last_modified_by"] = core_props.last_modified_by
        if core_props.revision is not None:
            metadata["revision"] = core_props.revision

        # Fallback: use filename as title if not in metadata
        if "title" not in metadata:
            metadata["title"] = path.stem

        # Extract paragraphs
        paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        metadata["paragraph_count"] = len(paragraphs)
        content_parts.extend(paragraphs)

        # Extract tables
        table_count = 0
        for table in doc.tables:
            table_count += 1
            table_text = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                if any(row_cells):
                    table_text.append(" | ".join(row_cells))
            if table_text:
                content_parts.append("\n[TABLE]\n" + "\n".join(table_text) + "\n[/TABLE]")

        metadata["table_count"] = table_count

        # Extract headers and footers
        header_text = []
        footer_text = []
        for section in doc.sections:
            if section.header and section.header.paragraphs:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        header_text.append(para.text.strip())
            if section.footer and section.footer.paragraphs:
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        footer_text.append(para.text.strip())

        if header_text:
            metadata["has_headers"] = True
        if footer_text:
            metadata["has_footers"] = True

    except (PackageNotFoundError, zipfile.BadZipFile) as e:
        logger.error(f"Invalid or corrupted DOCX file: {path} - {e}")
        raise DocumentParseError(f"Corrupted DOCX: {e}", path)
    except Exception as e:
        logger.error(f"Failed to read DOCX: {path} - {e}")
        raise DocumentParseError(f"DOCX read error: {e}", path)

    return "\n\n".join(content_parts), metadata


def _parse_yaml_frontmatter(content: str) -> Tuple[Dict[str, Any], str]:
    """Parse YAML front matter from markdown content.

    Returns:
        Tuple of (frontmatter_dict, content_without_frontmatter)
    """
    if not content.startswith("---"):
        return {}, content

    # Find the closing ---
    end_marker = content.find("\n---", 3)
    if end_marker == -1:
        return {}, content

    frontmatter_str = content[4:end_marker].strip()
    remaining_content = content[end_marker + 4:].lstrip("\n")

    # Parse YAML manually (simple key: value pairs)
    frontmatter = {}
    for line in frontmatter_str.split("\n"):
        line = line.strip()
        if not line or line.startswith("#"):
            continue

        if ":" in line:
            key, _, value = line.partition(":")
            key = key.strip().lower()
            value = value.strip()

            # Remove quotes if present
            if (value.startswith('"') and value.endswith('"')) or \
               (value.startswith("'") and value.endswith("'")):
                value = value[1:-1]

            # Handle arrays (simple format: [item1, item2])
            if value.startswith("[") and value.endswith("]"):
                items = value[1:-1].split(",")
                value = [item.strip().strip("\"'") for item in items if item.strip()]

            if value:
                frontmatter[key] = value

    return frontmatter, remaining_content


def load_markdown(path: Path) -> Tuple[str, Dict[str, Any]]:
    """Load Markdown/text file with structure-aware parsing and error handling."""
    metadata = {}

    try:
        content = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        # Try fallback encodings (latin-1 last as it accepts any byte sequence)
        for encoding in ["cp1252", "iso-8859-1", "latin-1"]:
            try:
                content = path.read_text(encoding=encoding)
                metadata["encoding_fallback"] = encoding
                logger.warning(f"Used fallback encoding {encoding} for: {path}")
                break
            except UnicodeDecodeError:
                continue
        else:
            logger.error(f"Failed to decode file with any encoding: {path}")
            raise DocumentParseError(f"Unable to decode file", path)
    except Exception as e:
        logger.error(f"Failed to read file: {path} - {e}")
        raise DocumentParseError(f"File read error: {e}", path)

    # Parse front matter if present
    frontmatter, _ = _parse_yaml_frontmatter(content)
    if frontmatter:
        metadata["has_frontmatter"] = True
        # Extract common frontmatter fields
        if "title" in frontmatter:
            metadata["title"] = frontmatter["title"]
        if "author" in frontmatter:
            metadata["author"] = frontmatter["author"]
        if "date" in frontmatter:
            metadata["creation_date"] = frontmatter["date"]
        if "tags" in frontmatter:
            metadata["tags"] = frontmatter["tags"]
        if "description" in frontmatter:
            metadata["description"] = frontmatter["description"]
        if "categories" in frontmatter:
            metadata["categories"] = frontmatter["categories"]

    # Extract markdown structure metadata
    lines = content.split("\n")

    # Count headers by level and extract first H1 for title fallback
    headers = {"h1": 0, "h2": 0, "h3": 0, "h4": 0, "h5": 0, "h6": 0}
    first_h1 = None
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("###### ") or stripped == "######":
            headers["h6"] += 1
        elif stripped.startswith("##### ") or stripped == "#####":
            headers["h5"] += 1
        elif stripped.startswith("#### ") or stripped == "####":
            headers["h4"] += 1
        elif stripped.startswith("### ") or stripped == "###":
            headers["h3"] += 1
        elif stripped.startswith("## ") or stripped == "##":
            headers["h2"] += 1
        elif stripped.startswith("# "):
            headers["h1"] += 1
            if first_h1 is None:
                first_h1 = stripped[2:].strip()
        elif stripped == "#":
            headers["h1"] += 1

    if any(headers.values()):
        metadata["headers"] = {k: v for k, v in headers.items() if v > 0}

    # Fallback: use first H1 as title, then filename
    if "title" not in metadata:
        if first_h1:
            metadata["title"] = first_h1
        else:
            metadata["title"] = path.stem

    # Detect code blocks
    code_block_count = content.count("```") // 2
    if code_block_count > 0:
        metadata["code_blocks"] = code_block_count

    # Detect lists (unordered: -, *, + and ordered: 1., 2., etc.)
    list_items = sum(1 for line in lines if _LIST_PATTERN.match(line))
    if list_items > 0:
        metadata["list_items"] = list_items

    # Detect links
    links = re.findall(r"\[([^\]]+)\]\(([^)]+)\)", content)
    if links:
        metadata["link_count"] = len(links)

    metadata["line_count"] = len(lines)

    return content, metadata


HANDLERS: Dict[str, Callable[[Path], Tuple[str, Dict[str, Any]]]] = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".md": load_markdown,
    ".txt": load_markdown,
}


@dataclass
class DocumentLoader:
    """Document loader with optional text normalization.

    Attributes:
        input_root: Root directory for input documents.
        normalizer: Optional TextNormalizer for additional text cleaning.
    """

    input_root: Path
    normalizer: Optional[TextNormalizer] = field(default=None)

    def load(self, path: Path) -> Tuple[Document, str]:
        """Load and process a document from disk.

        Args:
            path: Path to the document file.

        Returns:
            Tuple of (Document, content_hash).

        Raises:
            UnsupportedDocumentError: If no handler exists for the file type.
        """
        ext = path.suffix.lower()
        if ext not in HANDLERS:
            raise UnsupportedDocumentError(f"No loader configured for {ext} files.")
        handler = HANDLERS[ext]
        raw_text, extra_metadata = handler(path)

        # Apply basic text normalization
        normalized_text = normalize_text(raw_text)

        # Apply advanced normalization if configured
        normalization_result: Optional[NormalizationResult] = None
        if self.normalizer is not None:
            normalization_result = self.normalizer.normalize(normalized_text)
            normalized_text = normalization_result.text

        try:
            rel_path = str(path.relative_to(self.input_root)).replace("\\", "/")
        except ValueError:
            rel_path = path.name
        stat_info = path.stat()
        metadata = {
            "source_path": str(path.resolve()),
            "relative_path": rel_path,
            "file_extension": ext.lstrip("."),
            "last_modified": datetime.fromtimestamp(stat_info.st_mtime).isoformat(),
            "size_bytes": stat_info.st_size,
        }
        metadata.update(extra_metadata or {})
        content_hash = hash_file(path)
        metadata["content_hash"] = content_hash
        metadata["ingestion_timestamp"] = datetime.utcnow().isoformat() + "Z"

        # Add normalization metadata if normalization was applied
        if normalization_result is not None:
            metadata["normalization"] = {
                "original_length": normalization_result.original_length,
                "normalized_length": normalization_result.normalized_length,
                "rules_applied": normalization_result.rules_applied,
                "removed_patterns": normalization_result.removed_patterns,
            }

        doc_id = slugify(rel_path)
        document = Document(
            doc_id=doc_id,
            path=path,
            source_type=metadata["file_extension"],
            text=normalized_text,
            metadata=metadata,
        )
        return document, content_hash


def discover_documents(input_root: Path):
    documents = []
    for ext in HANDLERS:
        documents.extend(input_root.rglob(f"*{ext}"))
    documents.sort()
    return tuple(documents)
