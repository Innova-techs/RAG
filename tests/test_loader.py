"""Unit tests for document loaders (PDF, DOCX, Markdown)."""
from __future__ import annotations

from datetime import datetime
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

from ingestion.loader import (
    DocumentParseError,
    _parse_pdf_date,
    _parse_yaml_frontmatter,
    load_docx,
    load_markdown,
    load_pdf,
)


class TestPDFLoader:
    """Tests for PDF parser functionality."""

    def test_load_pdf_single_page(self, tmp_path: Path):
        """Test PDF with single page extracts text correctly."""
        with patch("PyPDF2.PdfReader") as mock_reader:
            mock_page = MagicMock()
            mock_page.extract_text.return_value = "Hello World"
            mock_reader.return_value.pages = [mock_page]
            mock_reader.return_value.is_encrypted = False

            pdf_path = tmp_path / "test.pdf"
            pdf_path.touch()

            text, metadata = load_pdf(pdf_path)

            assert text == "Hello World"
            assert metadata["page_count"] == 1
            assert metadata["is_encrypted"] is False

    def test_load_pdf_multi_page(self, tmp_path: Path):
        """Test PDF with multiple pages extracts all text."""
        with patch("PyPDF2.PdfReader") as mock_reader:
            mock_pages = []
            for i in range(3):
                mock_page = MagicMock()
                mock_page.extract_text.return_value = f"Page {i + 1} content"
                mock_pages.append(mock_page)

            mock_reader.return_value.pages = mock_pages
            mock_reader.return_value.is_encrypted = False

            pdf_path = tmp_path / "test.pdf"
            pdf_path.touch()

            text, metadata = load_pdf(pdf_path)

            assert "Page 1 content" in text
            assert "Page 2 content" in text
            assert "Page 3 content" in text
            assert metadata["page_count"] == 3

    def test_load_pdf_encrypted(self, tmp_path: Path):
        """Test encrypted PDF returns empty content with warning."""
        with patch("PyPDF2.PdfReader") as mock_reader:
            mock_reader.return_value.pages = []
            mock_reader.return_value.is_encrypted = True

            pdf_path = tmp_path / "encrypted.pdf"
            pdf_path.touch()

            text, metadata = load_pdf(pdf_path)

            assert text == ""
            assert metadata["is_encrypted"] is True
            assert metadata["parse_warning"] == "encrypted_pdf"

    def test_load_pdf_partial_extraction_failure(self, tmp_path: Path):
        """Test PDF with some pages failing extraction handles gracefully."""
        with patch("PyPDF2.PdfReader") as mock_reader:
            mock_page1 = MagicMock()
            mock_page1.extract_text.return_value = "Good page"

            mock_page2 = MagicMock()
            mock_page2.extract_text.side_effect = Exception("Extraction failed")

            mock_page3 = MagicMock()
            mock_page3.extract_text.return_value = "Another good page"

            mock_reader.return_value.pages = [mock_page1, mock_page2, mock_page3]
            mock_reader.return_value.is_encrypted = False

            pdf_path = tmp_path / "partial.pdf"
            pdf_path.touch()

            text, metadata = load_pdf(pdf_path)

            assert "Good page" in text
            assert "Another good page" in text
            assert metadata["failed_pages"] == [2]
            assert "partial_extraction" in metadata.get("parse_warning", "")

    def test_load_pdf_corrupted_raises_error(self, tmp_path: Path):
        """Test corrupted PDF raises DocumentParseError."""
        from PyPDF2.errors import PdfReadError

        with patch("PyPDF2.PdfReader") as mock_reader:
            mock_reader.side_effect = PdfReadError("Corrupted PDF")

            pdf_path = tmp_path / "corrupted.pdf"
            pdf_path.touch()

            with pytest.raises(DocumentParseError) as exc_info:
                load_pdf(pdf_path)

            assert "Corrupted PDF" in str(exc_info.value)

    def test_load_pdf_empty_pages(self, tmp_path: Path):
        """Test PDF with empty pages handles None text."""
        with patch("PyPDF2.PdfReader") as mock_reader:
            mock_page = MagicMock()
            mock_page.extract_text.return_value = None

            mock_reader.return_value.pages = [mock_page]
            mock_reader.return_value.is_encrypted = False

            pdf_path = tmp_path / "empty.pdf"
            pdf_path.touch()

            text, metadata = load_pdf(pdf_path)

            assert text == ""
            assert metadata["page_count"] == 1


class TestDOCXLoader:
    """Tests for DOCX parser functionality."""

    def test_load_docx_paragraphs(self, tmp_path: Path):
        """Test DOCX with paragraphs extracts text correctly."""
        with patch("docx.Document") as mock_doc:
            mock_para1 = MagicMock()
            mock_para1.text = "First paragraph"
            mock_para2 = MagicMock()
            mock_para2.text = "Second paragraph"

            mock_doc.return_value.paragraphs = [mock_para1, mock_para2]
            mock_doc.return_value.tables = []
            mock_doc.return_value.sections = []

            docx_path = tmp_path / "test.docx"
            docx_path.touch()

            text, metadata = load_docx(docx_path)

            assert "First paragraph" in text
            assert "Second paragraph" in text
            assert metadata["paragraph_count"] == 2

    def test_load_docx_with_tables(self, tmp_path: Path):
        """Test DOCX with tables extracts table content."""
        with patch("docx.Document") as mock_doc:
            mock_doc.return_value.paragraphs = []

            # Mock table with rows and cells
            mock_cell1 = MagicMock()
            mock_cell1.text = "Header 1"
            mock_cell2 = MagicMock()
            mock_cell2.text = "Header 2"

            mock_cell3 = MagicMock()
            mock_cell3.text = "Data 1"
            mock_cell4 = MagicMock()
            mock_cell4.text = "Data 2"

            mock_row1 = MagicMock()
            mock_row1.cells = [mock_cell1, mock_cell2]
            mock_row2 = MagicMock()
            mock_row2.cells = [mock_cell3, mock_cell4]

            mock_table = MagicMock()
            mock_table.rows = [mock_row1, mock_row2]

            mock_doc.return_value.tables = [mock_table]
            mock_doc.return_value.sections = []

            docx_path = tmp_path / "tables.docx"
            docx_path.touch()

            text, metadata = load_docx(docx_path)

            assert "[TABLE]" in text
            assert "Header 1 | Header 2" in text
            assert "Data 1 | Data 2" in text
            assert metadata["table_count"] == 1

    def test_load_docx_empty_paragraphs_skipped(self, tmp_path: Path):
        """Test DOCX skips empty paragraphs."""
        with patch("docx.Document") as mock_doc:
            mock_para1 = MagicMock()
            mock_para1.text = "Content"
            mock_para2 = MagicMock()
            mock_para2.text = "   "  # Whitespace only
            mock_para3 = MagicMock()
            mock_para3.text = ""  # Empty

            mock_doc.return_value.paragraphs = [mock_para1, mock_para2, mock_para3]
            mock_doc.return_value.tables = []
            mock_doc.return_value.sections = []

            docx_path = tmp_path / "sparse.docx"
            docx_path.touch()

            text, metadata = load_docx(docx_path)

            assert text == "Content"
            assert metadata["paragraph_count"] == 1

    def test_load_docx_corrupted_raises_error(self, tmp_path: Path):
        """Test corrupted DOCX raises DocumentParseError."""
        import zipfile

        with patch("docx.Document") as mock_doc:
            mock_doc.side_effect = zipfile.BadZipFile("Not a zip file")

            docx_path = tmp_path / "corrupted.docx"
            docx_path.touch()

            with pytest.raises(DocumentParseError) as exc_info:
                load_docx(docx_path)

            assert "Corrupted DOCX" in str(exc_info.value)

    def test_load_docx_with_headers_footers(self, tmp_path: Path):
        """Test DOCX with headers and footers sets metadata flags."""
        with patch("docx.Document") as mock_doc:
            mock_doc.return_value.paragraphs = []
            mock_doc.return_value.tables = []

            # Mock section with header
            mock_header_para = MagicMock()
            mock_header_para.text = "Document Header"
            mock_header = MagicMock()
            mock_header.paragraphs = [mock_header_para]

            mock_footer_para = MagicMock()
            mock_footer_para.text = "Page 1"
            mock_footer = MagicMock()
            mock_footer.paragraphs = [mock_footer_para]

            mock_section = MagicMock()
            mock_section.header = mock_header
            mock_section.footer = mock_footer

            mock_doc.return_value.sections = [mock_section]

            docx_path = tmp_path / "headers.docx"
            docx_path.touch()

            text, metadata = load_docx(docx_path)

            assert metadata.get("has_headers") is True
            assert metadata.get("has_footers") is True


class TestMarkdownLoader:
    """Tests for Markdown parser functionality."""

    def test_load_markdown_plain_text(self, tmp_path: Path):
        """Test plain text markdown loads correctly."""
        md_path = tmp_path / "plain.md"
        md_path.write_text("Just some plain text content.", encoding="utf-8")

        text, metadata = load_markdown(md_path)

        assert text == "Just some plain text content."
        assert metadata["line_count"] == 1

    def test_load_markdown_headers_detected(self, tmp_path: Path):
        """Test markdown headers are counted in metadata."""
        content = """# Main Title
## Section 1
### Subsection 1.1
## Section 2
### Subsection 2.1
### Subsection 2.2
"""
        md_path = tmp_path / "headers.md"
        md_path.write_text(content, encoding="utf-8")

        text, metadata = load_markdown(md_path)

        assert metadata["headers"]["h1"] == 1
        assert metadata["headers"]["h2"] == 2
        assert metadata["headers"]["h3"] == 3

    def test_load_markdown_hashtags_not_counted_as_headers(self, tmp_path: Path):
        """Test that hashtags (e.g., #hashtag) are not counted as headers."""
        content = """# Real Header
This is a #hashtag and another #tag in text.
##notaheader because no space
Check out #trending topics.
## Real H2 Header
"""
        md_path = tmp_path / "hashtags.md"
        md_path.write_text(content, encoding="utf-8")

        text, metadata = load_markdown(md_path)

        # Only real headers with space after # should be counted
        assert metadata["headers"]["h1"] == 1
        assert metadata["headers"]["h2"] == 1
        assert "h3" not in metadata["headers"]

    def test_load_markdown_code_blocks_detected(self, tmp_path: Path):
        """Test code blocks are counted in metadata."""
        content = """Some text

```python
def hello():
    print("Hello")
```

More text

```javascript
console.log("World");
```
"""
        md_path = tmp_path / "code.md"
        md_path.write_text(content, encoding="utf-8")

        text, metadata = load_markdown(md_path)

        assert metadata["code_blocks"] == 2

    def test_load_markdown_lists_detected(self, tmp_path: Path):
        """Test list items are counted in metadata."""
        content = """Shopping list:
- Apples
- Bananas
* Oranges

Steps:
1. First step
2. Second step
3. Third step
"""
        md_path = tmp_path / "lists.md"
        md_path.write_text(content, encoding="utf-8")

        text, metadata = load_markdown(md_path)

        assert metadata["list_items"] == 6

    def test_load_markdown_links_detected(self, tmp_path: Path):
        """Test links are counted in metadata."""
        content = """Check out [Google](https://google.com) and [GitHub](https://github.com)."""
        md_path = tmp_path / "links.md"
        md_path.write_text(content, encoding="utf-8")

        text, metadata = load_markdown(md_path)

        assert metadata["link_count"] == 2

    def test_load_markdown_frontmatter_detected(self, tmp_path: Path):
        """Test YAML frontmatter is detected."""
        content = """---
title: My Document
author: John Doe
---

# Content starts here
"""
        md_path = tmp_path / "frontmatter.md"
        md_path.write_text(content, encoding="utf-8")

        text, metadata = load_markdown(md_path)

        assert metadata.get("has_frontmatter") is True

    def test_load_markdown_encoding_fallback(self, tmp_path: Path):
        """Test fallback encoding for non-UTF-8 files."""
        md_path = tmp_path / "latin1.md"
        # Write with latin-1 encoding (cp1252 is tried first and will succeed)
        md_path.write_bytes("Café résumé naïve".encode("latin-1"))

        text, metadata = load_markdown(md_path)

        assert "Café" in text
        # cp1252 is tried first and succeeds for latin-1 compatible bytes
        assert metadata.get("encoding_fallback") == "cp1252"

    def test_load_markdown_encoding_fallback_cp1252_euro(self, tmp_path: Path):
        """Test cp1252 encoding with euro symbol (byte 0x80)."""
        md_path = tmp_path / "euro.md"
        # Euro symbol is 0x80 in cp1252, which is a control char in iso-8859-1
        # This verifies cp1252 is tried first and handles Windows-specific chars
        md_path.write_bytes(b"Price: \x80100 (euro symbol)")

        text, metadata = load_markdown(md_path)

        assert "100" in text
        assert metadata.get("encoding_fallback") == "cp1252"

    def test_load_markdown_lists_with_high_numbers(self, tmp_path: Path):
        """Test list detection with multi-digit ordered list numbers."""
        content = """1. First item
10. Tenth item
99. Ninety-ninth item
100. Hundredth item
999. Large number item
"""
        md_path = tmp_path / "high_numbers.md"
        md_path.write_text(content, encoding="utf-8")

        text, metadata = load_markdown(md_path)

        # All ordered list items should be detected regardless of number size
        assert metadata["list_items"] == 5

    def test_load_markdown_file_not_found(self, tmp_path: Path):
        """Test missing markdown file raises DocumentParseError."""
        md_path = tmp_path / "missing.md"

        with pytest.raises(DocumentParseError) as exc_info:
            load_markdown(md_path)

        assert "File read error" in str(exc_info.value)


class TestErrorHandling:
    """Tests for error handling across all loaders."""

    def test_pdf_general_exception_handling(self, tmp_path: Path):
        """Test general exceptions in PDF loading are caught."""
        with patch("PyPDF2.PdfReader") as mock_reader:
            mock_reader.side_effect = RuntimeError("Unexpected error")

            pdf_path = tmp_path / "error.pdf"
            pdf_path.touch()

            with pytest.raises(DocumentParseError) as exc_info:
                load_pdf(pdf_path)

            assert "PDF read error" in str(exc_info.value)

    def test_docx_package_not_found_error(self, tmp_path: Path):
        """Test PackageNotFoundError in DOCX loading is caught."""
        from docx.opc.exceptions import PackageNotFoundError

        with patch("docx.Document") as mock_doc:
            mock_doc.side_effect = PackageNotFoundError("Package not found")

            docx_path = tmp_path / "notfound.docx"
            docx_path.touch()

            with pytest.raises(DocumentParseError) as exc_info:
                load_docx(docx_path)

            assert "Corrupted DOCX" in str(exc_info.value)


class TestPDFMetadataExtraction:
    """Tests for PDF metadata extraction."""

    def test_parse_pdf_date_full_format(self):
        """Test parsing full PDF date format."""
        result = _parse_pdf_date("D:20240115103045")
        assert result == "2024-01-15T10:30:45"

    def test_parse_pdf_date_short_format(self):
        """Test parsing short PDF date format."""
        result = _parse_pdf_date("D:20240115")
        assert result == "2024-01-15"

    def test_parse_pdf_date_without_prefix(self):
        """Test parsing date without D: prefix."""
        result = _parse_pdf_date("20240115103045")
        assert result == "2024-01-15T10:30:45"

    def test_parse_pdf_date_none(self):
        """Test parsing None date."""
        result = _parse_pdf_date(None)
        assert result is None

    def test_parse_pdf_date_empty(self):
        """Test parsing empty date."""
        result = _parse_pdf_date("")
        assert result is None

    def test_load_pdf_extracts_metadata(self, tmp_path: Path):
        """Test PDF metadata extraction from document properties."""
        with patch("PyPDF2.PdfReader") as mock_reader:
            mock_page = MagicMock()
            mock_page.extract_text.return_value = "Content"

            mock_metadata = MagicMock()
            mock_metadata.title = "Test Document"
            mock_metadata.author = "John Doe"
            mock_metadata.creator = "Microsoft Word"
            mock_metadata.producer = "Adobe PDF"
            mock_metadata.subject = "Testing"
            mock_metadata.creation_date = "D:20240115103000"
            mock_metadata.modification_date = "D:20240120140000"

            mock_reader.return_value.pages = [mock_page]
            mock_reader.return_value.is_encrypted = False
            mock_reader.return_value.metadata = mock_metadata

            pdf_path = tmp_path / "test.pdf"
            pdf_path.touch()

            text, metadata = load_pdf(pdf_path)

            assert metadata["title"] == "Test Document"
            assert metadata["author"] == "John Doe"
            assert metadata["creator"] == "Microsoft Word"
            assert metadata["producer"] == "Adobe PDF"
            assert metadata["subject"] == "Testing"
            assert metadata["creation_date"] == "2024-01-15T10:30:00"
            assert metadata["modification_date"] == "2024-01-20T14:00:00"

    def test_load_pdf_title_fallback_to_filename(self, tmp_path: Path):
        """Test PDF uses filename as title when metadata is empty."""
        with patch("PyPDF2.PdfReader") as mock_reader:
            mock_page = MagicMock()
            mock_page.extract_text.return_value = "Content"

            mock_metadata = MagicMock()
            mock_metadata.title = None
            mock_metadata.author = None
            mock_metadata.creator = None
            mock_metadata.producer = None
            mock_metadata.subject = None
            mock_metadata.creation_date = None
            mock_metadata.modification_date = None

            mock_reader.return_value.pages = [mock_page]
            mock_reader.return_value.is_encrypted = False
            mock_reader.return_value.metadata = mock_metadata

            pdf_path = tmp_path / "my_document.pdf"
            pdf_path.touch()

            text, metadata = load_pdf(pdf_path)

            assert metadata["title"] == "my_document"


class TestDOCXMetadataExtraction:
    """Tests for DOCX metadata extraction."""

    def test_load_docx_extracts_core_properties(self, tmp_path: Path):
        """Test DOCX metadata extraction from core properties."""
        with patch("docx.Document") as mock_doc:
            mock_doc.return_value.paragraphs = []
            mock_doc.return_value.tables = []
            mock_doc.return_value.sections = []

            mock_core_props = MagicMock()
            mock_core_props.title = "Test Document"
            mock_core_props.author = "Jane Smith"
            mock_core_props.subject = "Testing"
            mock_core_props.keywords = "test, unit, python"
            mock_core_props.category = "Documentation"
            mock_core_props.comments = "Test file"
            mock_core_props.created = datetime(2024, 1, 15, 10, 30, 0)
            mock_core_props.modified = datetime(2024, 1, 20, 14, 0, 0)
            mock_core_props.last_modified_by = "John Doe"
            mock_core_props.revision = 5

            mock_doc.return_value.core_properties = mock_core_props

            docx_path = tmp_path / "test.docx"
            docx_path.touch()

            text, metadata = load_docx(docx_path)

            assert metadata["title"] == "Test Document"
            assert metadata["author"] == "Jane Smith"
            assert metadata["subject"] == "Testing"
            assert metadata["keywords"] == "test, unit, python"
            assert metadata["category"] == "Documentation"
            assert metadata["comments"] == "Test file"
            assert "2024-01-15" in metadata["creation_date"]
            assert "2024-01-20" in metadata["modification_date"]
            assert metadata["last_modified_by"] == "John Doe"
            assert metadata["revision"] == 5

    def test_load_docx_title_fallback_to_filename(self, tmp_path: Path):
        """Test DOCX uses filename as title when core properties empty."""
        with patch("docx.Document") as mock_doc:
            mock_doc.return_value.paragraphs = []
            mock_doc.return_value.tables = []
            mock_doc.return_value.sections = []

            mock_core_props = MagicMock()
            mock_core_props.title = None
            mock_core_props.author = None
            mock_core_props.subject = None
            mock_core_props.keywords = None
            mock_core_props.category = None
            mock_core_props.comments = None
            mock_core_props.created = None
            mock_core_props.modified = None
            mock_core_props.last_modified_by = None
            mock_core_props.revision = None

            mock_doc.return_value.core_properties = mock_core_props

            docx_path = tmp_path / "my_report.docx"
            docx_path.touch()

            text, metadata = load_docx(docx_path)

            assert metadata["title"] == "my_report"


class TestMarkdownMetadataExtraction:
    """Tests for Markdown metadata extraction."""

    def test_parse_yaml_frontmatter_basic(self):
        """Test parsing basic YAML front matter."""
        content = """---
title: My Document
author: John Doe
date: 2024-01-15
---

# Content"""
        frontmatter, remaining = _parse_yaml_frontmatter(content)

        assert frontmatter["title"] == "My Document"
        assert frontmatter["author"] == "John Doe"
        assert frontmatter["date"] == "2024-01-15"
        assert remaining.strip().startswith("# Content")

    def test_parse_yaml_frontmatter_with_quotes(self):
        """Test parsing front matter with quoted values."""
        content = """---
title: "My Document: A Study"
author: 'Jane Doe'
---

Content"""
        frontmatter, _ = _parse_yaml_frontmatter(content)

        assert frontmatter["title"] == "My Document: A Study"
        assert frontmatter["author"] == "Jane Doe"

    def test_parse_yaml_frontmatter_with_array(self):
        """Test parsing front matter with array values."""
        content = """---
tags: [python, testing, unit]
---

Content"""
        frontmatter, _ = _parse_yaml_frontmatter(content)

        assert frontmatter["tags"] == ["python", "testing", "unit"]

    def test_parse_yaml_frontmatter_no_frontmatter(self):
        """Test parsing content without front matter."""
        content = "# Just a heading\n\nSome content."
        frontmatter, remaining = _parse_yaml_frontmatter(content)

        assert frontmatter == {}
        assert remaining == content

    def test_load_markdown_extracts_frontmatter_metadata(self, tmp_path: Path):
        """Test markdown extracts metadata from front matter."""
        content = """---
title: Getting Started
author: Alice
date: 2024-06-01
tags: [tutorial, beginner]
description: A beginner's guide
---

# Introduction

Welcome to the guide."""
        md_path = tmp_path / "guide.md"
        md_path.write_text(content, encoding="utf-8")

        text, metadata = load_markdown(md_path)

        assert metadata["title"] == "Getting Started"
        assert metadata["author"] == "Alice"
        assert metadata["creation_date"] == "2024-06-01"
        assert metadata["tags"] == ["tutorial", "beginner"]
        assert metadata["description"] == "A beginner's guide"
        assert metadata["has_frontmatter"] is True

    def test_load_markdown_title_fallback_to_h1(self, tmp_path: Path):
        """Test markdown uses first H1 as title when no front matter."""
        content = """# My Amazing Article

Some introductory text.

## Section 1

More content."""
        md_path = tmp_path / "article.md"
        md_path.write_text(content, encoding="utf-8")

        text, metadata = load_markdown(md_path)

        assert metadata["title"] == "My Amazing Article"

    def test_load_markdown_title_fallback_to_filename(self, tmp_path: Path):
        """Test markdown uses filename as title when no H1 or front matter."""
        content = """Just some plain text without any headers."""
        md_path = tmp_path / "notes.md"
        md_path.write_text(content, encoding="utf-8")

        text, metadata = load_markdown(md_path)

        assert metadata["title"] == "notes"

    def test_load_markdown_frontmatter_overrides_h1(self, tmp_path: Path):
        """Test front matter title takes precedence over H1."""
        content = """---
title: Official Title
---

# Different H1 Title

Content here."""
        md_path = tmp_path / "test.md"
        md_path.write_text(content, encoding="utf-8")

        text, metadata = load_markdown(md_path)

        assert metadata["title"] == "Official Title"
